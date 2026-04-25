from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from src.analyzer.models import SEVERITY_ORDER, Finding
from src.core.models import MetadataSnapshot, ObjectInfo

if TYPE_CHECKING:
    from src.analyzer.engine import AnalyzerReport


@dataclass(slots=True)
class _AdviceItem:
    """Aggregated advice for a single rule across all findings.

    Sorting key: (severity rank, -occurrences) so the most severe and most
    frequent issues bubble up first.
    """

    rule_id: str
    title: str
    severity: str
    description: str
    rationale: str
    remediation: str
    targets: list[str]
    occurrences: int


# Translations local to the Word writer so it is self-contained: keeping the
# strings here avoids cluttering the UI translations dict and makes it easy
# to add new keys without touching the front end.
_LABELS: dict[str, dict[str, str]] = {
    "fr": {
        "data_dictionary_doc_title": "Data Dictionary",
        "data_dictionary_subtitle": "Data Dictionary a la date du {date}",
        "table_of_contents": "Table des matieres",
        "table_of_contents_hint": (
            "(Ouvrez ce document dans Word puis appuyez sur F9 ou cliquez "
            "droit \"Mettre a jour les champs\" pour rafraichir la table.)"
        ),
        "object_chapter_title": "{label} ({api_name})",
        "object_chapter_title_simple": "{api_name}",
        "section_information": "Informations",
        "section_fields": "Champs",
        "info_api_name": "API Name",
        "info_label": "Label",
        "info_plural_label": "Label pluriel",
        "info_custom": "Objet personnalise",
        "info_sharing_model": "Modele de partage",
        "info_deployment_status": "Statut de deploiement",
        "info_visibility": "Visibilite",
        "info_record_types": "Nombre de record types",
        "info_validation_rules": "Nombre de validation rules",
        "info_relationships": "Nombre de relations",
        "info_field_count": "Nombre de champs",
        "info_custom_field_count": "Nombre de champs custom",
        "info_description": "Description",
        "yes": "Oui",
        "no": "Non",
        "value_unspecified": "Non renseigne",
        "field_column_label": "Label",
        "field_column_api_name": "API Name",
        "field_column_type": "Type",
        "field_column_description": "Description",
        "field_no_description": "Aucune description fournie.",
        "no_objects": (
            "Aucun objet n'a ete documente : la liste de metadata est vide "
            "ou tous les objets sont presents dans le fichier d'exclusion."
        ),
        "summary_doc_title": "Resume de l'analyse",
        "summary_subtitle": "Resume genere le {date}",
        "section_overview": "Vue d'ensemble",
        "section_metrics": "Metriques de personnalisation",
        "section_findings": "Etat de l'analyse statique",
        "section_advice": "Conseils",
        "advice_intro": (
            "Les actions ci-dessous sont triees de la plus prioritaire a la "
            "moins prioritaire. La priorite combine la severite de la regle "
            "et le nombre d'occurrences detectees."
        ),
        "advice_no_findings": (
            "Aucune action critique a signaler : l'analyse statique n'a "
            "remonte aucun finding pour les regles activees."
        ),
        "advice_action": "Action {index} - {title}",
        "advice_severity": "Severite",
        "advice_occurrences": "Occurrences detectees",
        "advice_examples": "Exemples concernes",
        "advice_examples_more": "... et {count} autre(s).",
        "advice_description": "Constat",
        "advice_rationale": "Pourquoi c'est important",
        "advice_remediation": "Action recommandee",
        "overview_metrics_intro": (
            "Cette section presente les principaux indicateurs collectes "
            "lors de l'analyse de l'org."
        ),
        "metric_objects": "Objets analyses",
        "metric_custom_objects": "Objets personnalises",
        "metric_custom_fields": "Champs personnalises",
        "metric_record_types": "Record types",
        "metric_validation_rules": "Validation rules",
        "metric_layouts": "Page layouts",
        "metric_custom_tabs": "Onglets custom",
        "metric_custom_apps": "Applications custom",
        "metric_flows": "Flows",
        "metric_apex_classes": "Classes Apex",
        "metric_apex_triggers": "Triggers Apex",
        "metric_lwc": "Composants LWC",
        "metric_flexipages": "Pages Lightning (FlexiPages)",
        "metric_omni_scripts": "OmniScripts",
        "metric_omni_integration_procedures": "Integration Procedures",
        "metric_omni_ui_cards": "UI Cards / FlexCards",
        "metric_omni_data_transforms": "Data Transforms",
        "metric_score": "Score de personnalisation",
        "metric_level": "Niveau",
        "metric_adopt_adapt_score": "Score Adopt vs Adapt",
        "metric_adopt_adapt_level": "Niveau Adopt vs Adapt",
        "metric_findings_total": "Findings totaux",
        "metric_findings_critical": "Findings Critical",
        "metric_findings_major": "Findings Major",
        "metric_findings_minor": "Findings Minor",
        "metric_findings_info": "Findings Info",
        "overview_intro": (
            "Ce document presente une vue d'ensemble de l'org Salesforce "
            "apres l'analyse complete et la creation de la documentation. "
            "Il couvre les principales metriques, l'etat de l'analyse "
            "statique et les actions recommandees."
        ),
        "severity_critical": "Critique",
        "severity_major": "Majeur",
        "severity_minor": "Mineur",
        "severity_info": "Info",
    },
    "en": {
        "data_dictionary_doc_title": "Data Dictionary",
        "data_dictionary_subtitle": "Data Dictionary as of {date}",
        "table_of_contents": "Table of contents",
        "table_of_contents_hint": (
            "(Open this document in Word and press F9 or right-click "
            "\"Update Field\" to refresh the table.)"
        ),
        "object_chapter_title": "{label} ({api_name})",
        "object_chapter_title_simple": "{api_name}",
        "section_information": "Information",
        "section_fields": "Fields",
        "info_api_name": "API Name",
        "info_label": "Label",
        "info_plural_label": "Plural label",
        "info_custom": "Custom object",
        "info_sharing_model": "Sharing model",
        "info_deployment_status": "Deployment status",
        "info_visibility": "Visibility",
        "info_record_types": "Record type count",
        "info_validation_rules": "Validation rule count",
        "info_relationships": "Relationship count",
        "info_field_count": "Field count",
        "info_custom_field_count": "Custom field count",
        "info_description": "Description",
        "yes": "Yes",
        "no": "No",
        "value_unspecified": "Not specified",
        "field_column_label": "Label",
        "field_column_api_name": "API Name",
        "field_column_type": "Type",
        "field_column_description": "Description",
        "field_no_description": "No description provided.",
        "no_objects": (
            "No object has been documented: the metadata list is empty or "
            "every object is filtered by the exclusion file."
        ),
        "summary_doc_title": "Analysis summary",
        "summary_subtitle": "Summary generated on {date}",
        "section_overview": "Overview",
        "section_metrics": "Customization metrics",
        "section_findings": "Static analysis status",
        "section_advice": "Advice",
        "advice_intro": (
            "The actions below are ordered from highest to lowest priority. "
            "Priority combines the rule severity with the number of "
            "detected occurrences."
        ),
        "advice_no_findings": (
            "No critical action to flag: static analysis did not raise any "
            "finding for the enabled rules."
        ),
        "advice_action": "Action {index} - {title}",
        "advice_severity": "Severity",
        "advice_occurrences": "Detected occurrences",
        "advice_examples": "Affected items",
        "advice_examples_more": "... and {count} more.",
        "advice_description": "Finding",
        "advice_rationale": "Why it matters",
        "advice_remediation": "Recommended action",
        "overview_metrics_intro": (
            "This section presents the main indicators captured while "
            "analysing the org."
        ),
        "metric_objects": "Analysed objects",
        "metric_custom_objects": "Custom objects",
        "metric_custom_fields": "Custom fields",
        "metric_record_types": "Record types",
        "metric_validation_rules": "Validation rules",
        "metric_layouts": "Page layouts",
        "metric_custom_tabs": "Custom tabs",
        "metric_custom_apps": "Custom apps",
        "metric_flows": "Flows",
        "metric_apex_classes": "Apex classes",
        "metric_apex_triggers": "Apex triggers",
        "metric_lwc": "LWC components",
        "metric_flexipages": "Lightning pages (FlexiPages)",
        "metric_omni_scripts": "OmniScripts",
        "metric_omni_integration_procedures": "Integration Procedures",
        "metric_omni_ui_cards": "UI Cards / FlexCards",
        "metric_omni_data_transforms": "Data Transforms",
        "metric_score": "Customization score",
        "metric_level": "Level",
        "metric_adopt_adapt_score": "Adopt vs Adapt score",
        "metric_adopt_adapt_level": "Adopt vs Adapt level",
        "metric_findings_total": "Total findings",
        "metric_findings_critical": "Critical findings",
        "metric_findings_major": "Major findings",
        "metric_findings_minor": "Minor findings",
        "metric_findings_info": "Info findings",
        "overview_intro": (
            "This document provides an overview of the Salesforce org "
            "after the full analysis and documentation generation. It "
            "covers the main metrics, the static analysis status, and the "
            "recommended actions."
        ),
        "severity_critical": "Critical",
        "severity_major": "Major",
        "severity_minor": "Minor",
        "severity_info": "Info",
    },
}


# How many affected items we list per advice action - longer lists make the
# document hard to read and add little value once the user reproduces the
# pattern locally.
_ADVICE_TARGET_LIMIT = 8


class WordReportWriter:
    """Generates the Word counterparts of the documentation.

    The writer is intentionally decoupled from the rest of the reporting
    pipeline: it only consumes plain dataclasses (`MetadataSnapshot`,
    `AnalyzerReport`) and writes ``.docx`` files into the directory that
    callers specify.
    """

    def __init__(self, language: str = "fr", log_callback=None) -> None:
        self.language = language if language in _LABELS else "fr"
        self.log = log_callback or (lambda message: None)

    # ------------------------------------------------------------------ public API

    def write_data_dictionary_document(
        self,
        snapshot: MetadataSnapshot,
        output_path: str | Path,
    ) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Excluded objects are already filtered upstream in the parser; we
        # additionally drop objects that have no field at all so the
        # generated document mirrors the Excel dictionary exactly.
        documented_objects = [obj for obj in snapshot.objects if obj.fields]

        document = Document()
        self._configure_default_style(document)
        self._enable_field_auto_update(document)

        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
        self._add_cover_page(
            document,
            title=self._t("data_dictionary_doc_title"),
            subtitle=self._t("data_dictionary_subtitle", date=generated_at),
        )

        self._add_table_of_contents_page(document)

        if not documented_objects:
            document.add_paragraph(self._t("no_objects"))
        else:
            for index, obj in enumerate(documented_objects):
                if index > 0:
                    document.add_page_break()
                self._add_object_chapter(document, obj)

        document.save(output_path)
        self.log(
            f"Data Dictionary Word genere ({len(documented_objects)} objet(s)) : "
            f"{output_path}"
        )
        return output_path

    def write_summary_document(
        self,
        snapshot: MetadataSnapshot,
        analyzer_report: "AnalyzerReport | None",
        output_path: str | Path,
    ) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        self._configure_default_style(document)

        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
        self._add_cover_page(
            document,
            title=self._t("summary_doc_title"),
            subtitle=self._t("summary_subtitle", date=generated_at),
        )

        document.add_heading(self._t("section_overview"), level=1)
        document.add_paragraph(self._t("overview_intro"))

        document.add_heading(self._t("section_metrics"), level=1)
        document.add_paragraph(self._t("overview_metrics_intro"))
        self._add_metrics_table(document, snapshot, analyzer_report)

        document.add_heading(self._t("section_advice"), level=1)
        advice_items = self._build_advice_items(analyzer_report)
        if not advice_items:
            document.add_paragraph(self._t("advice_no_findings"))
        else:
            document.add_paragraph(self._t("advice_intro"))
            for index, advice in enumerate(advice_items, start=1):
                self._add_advice_section(document, advice, index)

        document.save(output_path)
        self.log(f"Resume Word genere : {output_path}")
        return output_path

    # ------------------------------------------------------------------ helpers

    def _t(self, key: str, **fmt: object) -> str:
        labels = _LABELS.get(self.language) or _LABELS["fr"]
        template = labels.get(key, key)
        if not fmt:
            return template
        try:
            return template.format(**fmt)
        except (KeyError, IndexError):
            return template

    @staticmethod
    def _configure_default_style(document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    @staticmethod
    def _enable_field_auto_update(document: Document) -> None:
        # Tells Word to prompt the user to update fields (i.e. the TOC) when
        # the document is opened. Without this the TOC stays empty until the
        # user manually presses F9.
        settings = document.settings.element
        existing = settings.find(qn("w:updateFields"))
        if existing is None:
            update_fields = OxmlElement("w:updateFields")
            update_fields.set(qn("w:val"), "true")
            settings.append(update_fields)

    def _add_cover_page(self, document: Document, *, title: str, subtitle: str) -> None:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = 1  # center
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(28)

        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = 1
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.font.size = Pt(16)

        # Force the cover to occupy a full page so the table of contents
        # naturally lands on page 2.
        document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    def _add_table_of_contents_page(self, document: Document) -> None:
        document.add_heading(self._t("table_of_contents"), level=1)
        hint = document.add_paragraph(self._t("table_of_contents_hint"))
        for run in hint.runs:
            run.italic = True

        toc_paragraph = document.add_paragraph()
        run = toc_paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        # h: hyperlinks, z: hide tab leader on web, u: use heading styles.
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_separate)
        run._r.append(fld_end)

        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_object_chapter(self, document: Document, obj: ObjectInfo) -> None:
        if obj.label:
            heading = self._t(
                "object_chapter_title",
                label=obj.label,
                api_name=obj.api_name,
            )
        else:
            heading = self._t("object_chapter_title_simple", api_name=obj.api_name)
        document.add_heading(heading, level=1)

        document.add_heading(self._t("section_information"), level=2)
        self._add_information_table(document, obj)

        document.add_heading(self._t("section_fields"), level=2)
        self._add_fields_table(document, obj)

    def _add_information_table(self, document: Document, obj: ObjectInfo) -> None:
        rows: list[tuple[str, str]] = [
            (self._t("info_api_name"), obj.api_name or self._t("value_unspecified")),
            (self._t("info_label"), obj.label or self._t("value_unspecified")),
            (
                self._t("info_plural_label"),
                obj.plural_label or self._t("value_unspecified"),
            ),
            (
                self._t("info_custom"),
                self._t("yes") if obj.custom else self._t("no"),
            ),
            (
                self._t("info_sharing_model"),
                obj.sharing_model or self._t("value_unspecified"),
            ),
            (
                self._t("info_deployment_status"),
                obj.deployment_status or self._t("value_unspecified"),
            ),
            (
                self._t("info_visibility"),
                obj.visibility or self._t("value_unspecified"),
            ),
            (self._t("info_field_count"), str(len(obj.fields))),
            (
                self._t("info_custom_field_count"),
                str(sum(1 for field in obj.fields if field.custom)),
            ),
            (self._t("info_record_types"), str(len(obj.record_types))),
            (self._t("info_validation_rules"), str(len(obj.validation_rules))),
            (self._t("info_relationships"), str(len(obj.relationships))),
            (
                self._t("info_description"),
                (obj.description or "").strip() or self._t("value_unspecified"),
            ),
        ]
        table = document.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for index, (label, value) in enumerate(rows):
            label_cell = table.cell(index, 0)
            value_cell = table.cell(index, 1)
            self._set_cell_text(label_cell, label, bold=True)
            self._set_cell_text(value_cell, value)
        self._set_table_column_widths(table, [Cm(5.5), Cm(11.0)])

    def _add_fields_table(self, document: Document, obj: ObjectInfo) -> None:
        headers = [
            self._t("field_column_label"),
            self._t("field_column_api_name"),
            self._t("field_column_type"),
            self._t("field_column_description"),
        ]
        table = document.add_table(rows=1 + len(obj.fields), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for col_index, header in enumerate(headers):
            self._set_cell_text(table.cell(0, col_index), header, bold=True)

        sorted_fields = sorted(
            obj.fields,
            key=lambda f: ((f.label or f.api_name or "").lower()),
        )
        for row_index, field_info in enumerate(sorted_fields, start=1):
            description = (field_info.description or "").strip()
            if not description:
                description = self._t("field_no_description")
            self._set_cell_text(
                table.cell(row_index, 0),
                field_info.label or field_info.api_name,
            )
            self._set_cell_text(table.cell(row_index, 1), field_info.api_name)
            self._set_cell_text(
                table.cell(row_index, 2),
                field_info.data_type or self._t("value_unspecified"),
            )
            self._set_cell_text(table.cell(row_index, 3), description)

        self._set_table_column_widths(
            table,
            [Cm(4.5), Cm(4.5), Cm(2.5), Cm(5.0)],
        )

    def _add_metrics_table(
        self,
        document: Document,
        snapshot: MetadataSnapshot,
        analyzer_report: "AnalyzerReport | None",
    ) -> None:
        metrics = snapshot.metrics
        rows: list[tuple[str, str]] = [
            (self._t("metric_objects"), str(len(snapshot.objects))),
            (self._t("metric_custom_objects"), str(metrics.custom_objects)),
            (self._t("metric_custom_fields"), str(metrics.custom_fields)),
            (self._t("metric_record_types"), str(metrics.record_types)),
            (self._t("metric_validation_rules"), str(metrics.validation_rules)),
            (self._t("metric_layouts"), str(metrics.layouts)),
            (self._t("metric_custom_tabs"), str(metrics.custom_tabs)),
            (self._t("metric_custom_apps"), str(metrics.custom_apps)),
            (self._t("metric_flows"), str(metrics.flows)),
            (self._t("metric_apex_classes"), str(metrics.apex_classes)),
            (self._t("metric_apex_triggers"), str(metrics.apex_triggers)),
            (self._t("metric_lwc"), str(metrics.lwc_count)),
            (self._t("metric_flexipages"), str(metrics.flexipage_count)),
            (self._t("metric_omni_scripts"), str(metrics.omni_scripts)),
            (
                self._t("metric_omni_integration_procedures"),
                str(metrics.omni_integration_procedures),
            ),
            (self._t("metric_omni_ui_cards"), str(metrics.omni_ui_cards)),
            (
                self._t("metric_omni_data_transforms"),
                str(metrics.omni_data_transforms),
            ),
            (self._t("metric_score"), str(metrics.score)),
            (self._t("metric_level"), metrics.level),
            (self._t("metric_adopt_adapt_score"), str(metrics.adopt_adapt_score)),
            (self._t("metric_adopt_adapt_level"), metrics.adopt_adapt_level),
        ]
        if analyzer_report is not None:
            severity_counts = analyzer_report.severity_counts()
            total = sum(severity_counts.values())
            rows.extend(
                [
                    (self._t("metric_findings_total"), str(total)),
                    (
                        self._t("metric_findings_critical"),
                        str(severity_counts.get("Critical", 0)),
                    ),
                    (
                        self._t("metric_findings_major"),
                        str(severity_counts.get("Major", 0)),
                    ),
                    (
                        self._t("metric_findings_minor"),
                        str(severity_counts.get("Minor", 0)),
                    ),
                    (
                        self._t("metric_findings_info"),
                        str(severity_counts.get("Info", 0)),
                    ),
                ]
            )

        table = document.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        for index, (label, value) in enumerate(rows):
            self._set_cell_text(table.cell(index, 0), label, bold=True)
            self._set_cell_text(table.cell(index, 1), value)
        self._set_table_column_widths(table, [Cm(7.0), Cm(8.0)])

    def _build_advice_items(
        self, analyzer_report: "AnalyzerReport | None"
    ) -> list[_AdviceItem]:
        if analyzer_report is None:
            return []

        findings_by_rule: dict[str, list[Finding]] = {}
        for finding in analyzer_report.all_findings():
            findings_by_rule.setdefault(finding.rule.id, []).append(finding)

        items: list[_AdviceItem] = []
        for rule_id, findings in findings_by_rule.items():
            rule = findings[0].rule
            target_counter: Counter[str] = Counter()
            for finding in findings:
                key = f"{finding.target_kind}: {finding.target_name}"
                target_counter[key] += 1
            ordered_targets = [
                target for target, _ in target_counter.most_common(_ADVICE_TARGET_LIMIT)
            ]
            items.append(
                _AdviceItem(
                    rule_id=rule_id,
                    title=rule.title or rule_id,
                    severity=rule.severity,
                    description=rule.description,
                    rationale=rule.rationale,
                    remediation=rule.remediation,
                    targets=ordered_targets,
                    occurrences=len(findings),
                )
            )

        items.sort(
            key=lambda item: (
                SEVERITY_ORDER.get(item.severity, 99),
                -item.occurrences,
                item.rule_id,
            )
        )
        return items

    def _add_advice_section(
        self, document: Document, advice: _AdviceItem, index: int
    ) -> None:
        document.add_heading(
            self._t("advice_action", index=index, title=advice.title),
            level=2,
        )

        meta_paragraph = document.add_paragraph()
        meta_paragraph.add_run(
            f"{self._t('advice_severity')}: "
        ).bold = True
        meta_paragraph.add_run(self._severity_label(advice.severity))
        meta_paragraph.add_run("    ")
        meta_paragraph.add_run(
            f"{self._t('advice_occurrences')}: "
        ).bold = True
        meta_paragraph.add_run(str(advice.occurrences))

        if advice.description:
            self._add_labelled_paragraph(
                document, self._t("advice_description"), advice.description
            )
        if advice.rationale:
            self._add_labelled_paragraph(
                document, self._t("advice_rationale"), advice.rationale
            )
        if advice.remediation:
            self._add_labelled_paragraph(
                document, self._t("advice_remediation"), advice.remediation
            )

        if advice.targets:
            document.add_paragraph(self._t("advice_examples"), style="Heading 3")
            for target in advice.targets:
                document.add_paragraph(target, style="List Bullet")
            remaining = advice.occurrences - len(advice.targets)
            if remaining > 0:
                document.add_paragraph(
                    self._t("advice_examples_more", count=remaining)
                ).runs[0].italic = True

    def _severity_label(self, severity: str) -> str:
        mapping = {
            "Critical": self._t("severity_critical"),
            "Major": self._t("severity_major"),
            "Minor": self._t("severity_minor"),
            "Info": self._t("severity_info"),
        }
        return mapping.get(severity, severity)

    @staticmethod
    def _add_labelled_paragraph(document: Document, label: str, body: str) -> None:
        cleaned = " ".join((body or "").split())
        if not cleaned:
            return
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(cleaned)

    @staticmethod
    def _set_cell_text(cell, value: str, *, bold: bool = False) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.bold = bold

    @staticmethod
    def _set_table_column_widths(table, widths) -> None:
        for row in table.rows:
            for index, width in enumerate(widths):
                if index < len(row.cells):
                    row.cells[index].width = width
